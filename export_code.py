#!/usr/bin/env python3
"""
Script d'export du code et de l'architecture du projet dans un fichier Word (.docx).

Utilisation :
    1. Installer la dépendance :
        pip install python-docx

    2. Placer ce fichier `export_code.py` à la racine du projet (là où se trouve package.json / vite.config.js).

    3. Lancer :
        python export_code.py

    4. Le fichier `export_code.docx` sera généré à la racine du projet.
"""

import os
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt


# ------------ PARAMÈTRES À ADAPTER AU BESOIN ------------ #

# Extensions considérées comme "fichiers de code"
CODE_EXTENSIONS = {
    ".py", ".js", ".jsx", ".ts", ".tsx",
    ".html", ".htm", ".css",
    ".json", ".yml", ".yaml", ".toml",
    ".md", ".txt",
    ".env", ".env.example", ".env.sample",
}

# Dossiers à ignorer pour l'export
IGNORED_DIRS = {
    "node_modules",
    ".git",
    ".idea",
    ".vscode",
    "dist",
    "build",
    "__pycache__",
    ".pytest_cache",
    ".venv",
    "venv",
    ".next",
    ".turbo",
}


# ------------ FONCTIONS UTILITAIRES ------------ #

def should_ignore_dir(dirname: str) -> bool:
    """Retourne True si le dossier doit être ignoré."""
    return dirname in IGNORED_DIRS


def is_code_file(path: Path) -> bool:
    """Retourne True si le fichier correspond à une extension de code connue."""
    if path.is_dir():
        return False
    ext = path.suffix.lower()
    if ext in CODE_EXTENSIONS:
        return True
    # Gestion spéciale pour les fichiers .env, .env.example, etc.
    name = path.name
    if name.startswith(".env"):
        return True
    return False


def build_tree(root: Path) -> str:
    """
    Construit une représentation en arbre du projet (type `tree`).
    """
    lines = []

    def _walk(current: Path, prefix: str = ""):
        entries = [e for e in current.iterdir()
                   if not (e.is_dir() and should_ignore_dir(e.name))]
        # Tri : dossiers d'abord, puis fichiers
        dirs = sorted([e for e in entries if e.is_dir()], key=lambda p: p.name.lower())
        files = sorted([e for e in entries if e.is_file()], key=lambda p: p.name.lower())
        children = dirs + files

        for index, entry in enumerate(children):
            connector = "└── " if index == len(children) - 1 else "├── "
            lines.append(f"{prefix}{connector}{entry.name}")

            if entry.is_dir():
                extension = "    " if index == len(children) - 1 else "│   "
                _walk(entry, prefix + extension)

    lines.append(root.name)
    _walk(root)
    return "\n".join(lines)


def add_code_paragraph(doc: Document, text: str):
    """
    Ajoute un paragraphe "type code" dans le document (police monospace).
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    font = run.font
    font.name = "Courier New"
    font.size = Pt(9)


# ------------ LOGIQUE PRINCIPALE ------------ #

def main():
    root_dir = Path(__file__).resolve().parent
    output_path = root_dir / "export_code.docx"

    document = Document()

    # Titre principal
    document.add_heading("Export du code du projet", level=0)
    document.add_paragraph(f"Racine du projet : {root_dir}")
    document.add_paragraph(f"Généré le : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph("")  # Ligne vide

    # Section : Architecture
    document.add_heading("1. Architecture du projet", level=1)
    tree_str = build_tree(root_dir)
    add_code_paragraph(document, tree_str)

    # Section : Fichiers de code
    document.add_page_break()
    document.add_heading("2. Fichiers de code", level=1)

    # Parcours de tous les fichiers
    for dirpath, dirnames, filenames in os.walk(root_dir):
        # Filtrer les dossiers à ignorer (in-place pour que os.walk les saute)
        dirnames[:] = [d for d in dirnames if not should_ignore_dir(d)]

        for filename in filenames:
            file_path = Path(dirpath) / filename

            # Ignorer nous-même si tu veux
            # if file_path == Path(__file__).resolve():
            #     continue

            if not is_code_file(file_path):
                continue

            rel_path = file_path.relative_to(root_dir)

            # Ajouter un sous-titre pour chaque fichier
            document.add_heading(str(rel_path), level=2)

            # Lire le contenu du fichier
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            except UnicodeDecodeError:
                # Si problème d'encodage, on retente en "latin-1" ou on ignore les erreurs
                try:
                    with open(file_path, "r", encoding="latin-1") as f:
                        content = f.read()
                except Exception:
                    content = "// Impossible de lire ce fichier (problème d'encodage)"

            # Ajouter le contenu
            add_code_paragraph(document, content)
            # Petite séparation entre fichiers
            document.add_paragraph("\n")

    # Sauvegarde du document
    document.save(output_path)
    print(f"✅ Fichier export généré : {output_path}")


if __name__ == "__main__":
    main()
